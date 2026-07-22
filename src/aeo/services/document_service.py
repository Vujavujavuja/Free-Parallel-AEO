"""Extract plain text from uploaded documents (docx, pdf, md, txt) so the
orchestrator can ground question generation in the company's own material."""

from __future__ import annotations

import io

from aeo.schemas.company import SourceDocument

MAX_CHARS = 20_000  # stored/parsed cap per document
SUPPORTED = {".docx", ".pdf", ".md", ".markdown", ".txt", ".text"}


def _ext(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot != -1 else ""


def extract_text(filename: str, data: bytes) -> str:
    ext = _ext(filename)
    if ext == ".docx":
        text = _from_docx(data)
    elif ext == ".pdf":
        text = _from_pdf(data)
    else:  # md / markdown / txt / unknown -> best-effort utf-8
        text = data.decode("utf-8", errors="replace")
    return text.strip()[:MAX_CHARS]


def parse_document(filename: str, data: bytes) -> SourceDocument:
    return SourceDocument(name=filename, text=extract_text(filename, data))


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)
