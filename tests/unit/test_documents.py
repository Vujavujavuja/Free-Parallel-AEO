"""Document parsing unit tests (docx/md/txt)."""

from __future__ import annotations

import io

from aeo.services.document_service import extract_text


def test_extract_markdown() -> None:
    text = extract_text("brief.md", b"# Title\nBody line with detail.")
    assert "Body line with detail." in text


def test_extract_txt() -> None:
    assert extract_text("notes.txt", b"plain text here") == "plain text here"


def test_extract_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph mentioning Acme.")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("brief.docx", buf.getvalue())
    assert "First paragraph." in text
    assert "Acme" in text


def test_truncation_cap() -> None:
    big = b"x" * 30_000
    assert len(extract_text("big.txt", big)) == 20_000
